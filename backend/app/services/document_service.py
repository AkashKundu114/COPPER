import csv
import datetime
import io
import json
import os
import re
from pathlib import Path
from typing import Any

from app.ai.memory.memory_manager import memory_manager
from app.core.config import settings
from app.core.document_indexer import chunk_text
from app.core.logger import logger


class DocumentService:
    """
    Robust Document Parser, Extractor, Creator, and Formatter Service for C.O.P.P.E.R.
    Supports parsing & generating PDF, DOCX (Word), Markdown, HTML, CSV/TSV, JSON, YAML, and Plain Text.
    """

    SUPPORTED_EXTENSIONS = {
        "pdf": "PDF Document (.pdf)",
        "docx": "Microsoft Word (.docx)",
        "txt": "Plain Text (.txt)",
        "md": "Markdown Document (.md)",
        "html": "HTML Web Document (.html)",
        "csv": "Comma-Separated Values (.csv)",
        "tsv": "Tab-Separated Values (.tsv)",
        "json": "JSON Structured Data (.json)",
        "yaml": "YAML Configuration (.yaml)",
        "yml": "YAML Configuration (.yml)",
        "py": "Python Source (.py)",
        "js": "JavaScript Source (.js)",
        "ts": "TypeScript Source (.ts)",
        "tsx": "React TypeScript Source (.tsx)",
        "jsx": "React JavaScript Source (.jsx)",
        "css": "CSS Stylesheet (.css)",
        "sql": "SQL Database Script (.sql)",
        "xml": "XML Data (.xml)",
        "log": "System Log (.log)",
        "sh": "Shell Script (.sh)",
        "bat": "Batch Script (.bat)",
        "ps1": "PowerShell Script (.ps1)",
        "rs": "Rust Source (.rs)",
        "go": "Go Source (.go)",
        "java": "Java Source (.java)",
        "cpp": "C++ Source (.cpp)",
    }

    DOCUMENT_TEMPLATES = {
        "technical_report": {
            "id": "technical_report",
            "name": "Technical Architecture & System Report",
            "description": "In-depth engineering report with executive summary, system architecture, component breakdown, and benchmarks.",
            "default_format": "pdf",
            "recommended_sections": ["Executive Summary", "System Architecture", "Component Specifications", "Security & Data Governance", "Performance Benchmarks", "Recommendations"],
        },
        "project_proposal": {
            "id": "project_proposal",
            "name": "Project Proposal & Implementation Plan",
            "description": "Formal proposal outlining scope, business justification, milestones, resource allocation, and risk matrix.",
            "default_format": "docx",
            "recommended_sections": ["Problem Statement", "Proposed Solution", "Project Scope & Deliverables", "Timeline & Milestones", "Resource Requirements", "Risk Analysis"],
        },
        "executive_summary": {
            "id": "executive_summary",
            "name": "Executive Briefing Document",
            "description": "High-impact summary tailored for leadership with key findings, metric highlights, and strategic decisions.",
            "default_format": "pdf",
            "recommended_sections": ["Strategic Context", "Key Findings & Metrics", "Impact Analysis", "Actionable Next Steps"],
        },
        "meeting_notes": {
            "id": "meeting_notes",
            "name": "Meeting Minutes & Action Items",
            "description": "Structured agenda, discussion points, decisions made, and assigned action item table.",
            "default_format": "md",
            "recommended_sections": ["Meeting Overview & Attendees", "Agenda Topics", "Discussion Summary", "Key Decisions Made", "Action Items & Owners"],
        },
        "resume": {
            "id": "resume",
            "name": "Professional Curriculum Vitae / Resume",
            "description": "Polished career profile with summary, core competencies, professional experience, education, and technical skills.",
            "default_format": "pdf",
            "recommended_sections": ["Professional Summary", "Core Competencies", "Work Experience", "Technical Skills", "Education & Certifications"],
        },
        "formal_letter": {
            "id": "formal_letter",
            "name": "Formal Business / Official Letter",
            "description": "Standard business letter format with header block, salutation, body paragraphs, call to action, and signature block.",
            "default_format": "docx",
            "recommended_sections": ["Recipient Information", "Subject Line", "Salutation", "Opening Context", "Core Message", "Call to Action", "Sign-off"],
        },
        "invoice_table": {
            "id": "invoice_table",
            "name": "Tabular Data & Financial Sheet",
            "description": "Structured columns, line items, unit costs, subtotals, taxes, and grand totals.",
            "default_format": "csv",
            "recommended_sections": ["Item Description", "Category", "Quantity", "Unit Rate", "Subtotal", "Tax", "Total"],
        },
        "research_paper": {
            "id": "research_paper",
            "name": "Academic & Scientific Research Paper",
            "description": "Formal research layout with abstract, literature review, methodology, findings, and bibliography.",
            "default_format": "html",
            "recommended_sections": ["Abstract", "Introduction", "Related Work", "Methodology", "Experimental Results", "Discussion", "References"],
        },
    }

    def __init__(self, output_dir: str | None = None):
        self.output_dir = Path(output_dir or getattr(settings, "DOCUMENTS_DIR", "data/documents"))
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def format_file_size(self, size_bytes: int) -> str:
        if size_bytes >= 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.2f} MB"
        if size_bytes >= 1024:
            return f"{size_bytes / 1024:.1f} KB"
        return f"{size_bytes} B"

    def _sanitize_filename(self, filename: str, ext: str) -> str:
        clean = re.sub(r"[^\w\s-]", "", filename).strip()
        clean = re.sub(r"[\s]+", "_", clean)
        if not clean:
            clean = f"document_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        if not clean.lower().endswith(f".{ext.lower()}"):
            clean = f"{clean}.{ext.lower()}"
        return clean

    # =========================================================================
    # DOCUMENT GENERATION ENGINE
    # =========================================================================

    def generate_pdf(
        self,
        title: str,
        sections: list[dict[str, Any]],
        filename: str | None = None,
        author: str = "C.O.P.P.E.R. AI",
        metadata: dict[str, Any] | None = None,
    ) -> tuple[bytes, Path]:
        """
        Generates a styled, multi-section PDF document using ReportLab.
        Falls back to PyMuPDF if needed.
        """
        fname = self._sanitize_filename(filename or title, "pdf")
        out_path = self.output_dir / fname

        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
                HRFlowable,
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=40,
                leftMargin=40,
                topMargin=40,
                bottomMargin=40,
            )

            styles = getSampleStyleSheet()

            # Custom Copper Styles
            copper_dark = colors.HexColor("#0f172a")
            copper_accent = colors.HexColor("#c87d55")
            copper_slate = colors.HexColor("#475569")
            copper_bg = colors.HexColor("#f8fafc")

            title_style = ParagraphStyle(
                "DocTitle",
                parent=styles["Normal"],
                fontName="Helvetica-Bold",
                fontSize=22,
                leading=26,
                textColor=copper_dark,
                alignment=TA_LEFT,
                spaceAfter=6,
            )

            subtitle_style = ParagraphStyle(
                "DocSubtitle",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=10,
                leading=14,
                textColor=copper_slate,
                spaceAfter=15,
            )

            h1_style = ParagraphStyle(
                "DocH1",
                parent=styles["Normal"],
                fontName="Helvetica-Bold",
                fontSize=14,
                leading=18,
                textColor=copper_accent,
                spaceBefore=14,
                spaceAfter=6,
            )

            h2_style = ParagraphStyle(
                "DocH2",
                parent=styles["Normal"],
                fontName="Helvetica-Bold",
                fontSize=11,
                leading=15,
                textColor=copper_dark,
                spaceBefore=8,
                spaceAfter=4,
            )

            body_style = ParagraphStyle(
                "DocBody",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=9.5,
                leading=14,
                textColor=copper_dark,
                spaceAfter=8,
            )

            bullet_style = ParagraphStyle(
                "DocBullet",
                parent=styles["Normal"],
                fontName="Helvetica",
                fontSize=9.5,
                leading=13,
                textColor=copper_dark,
                leftIndent=15,
                firstLineIndent=-10,
                spaceAfter=4,
            )

            story = []

            # Document Header / Title Block
            story.append(Paragraph(title, title_style))
            timestamp = datetime.datetime.now().strftime("%B %d, %Y · %H:%M UTC")
            meta_line = f"Generated by <b>{author}</b> · {timestamp}"
            story.append(Paragraph(meta_line, subtitle_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=copper_accent, spaceAfter=15))

            # Process Sections
            for sec in sections:
                heading = sec.get("heading") or sec.get("title")
                if heading:
                    story.append(Paragraph(heading, h1_style))

                subheading = sec.get("subheading")
                if subheading:
                    story.append(Paragraph(subheading, h2_style))

                content = sec.get("content") or sec.get("text") or ""
                if content:
                    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
                    for p in paragraphs:
                        clean_p = p.replace("\n", "<br/>")
                        story.append(Paragraph(clean_p, body_style))

                bullets = sec.get("bullets") or sec.get("items") or []
                if isinstance(bullets, list):
                    for b in bullets:
                        story.append(Paragraph(f"• {b}", bullet_style))

                # Handle Section Tables
                table_data = sec.get("table")
                if table_data and isinstance(table_data, dict):
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])
                    if headers or rows:
                        grid = [headers] + rows if headers else rows
                        table_flow = Table(grid, colWidths=None)
                        table_flow.setStyle(
                            TableStyle(
                                [
                                    ("BACKGROUND", (0, 0), (-1, 0), copper_accent),
                                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                                    ("BACKGROUND", (0, 1), (-1, -1), copper_bg),
                                    ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
                                ]
                            )
                        )
                        story.append(Spacer(1, 6))
                        story.append(table_flow)
                        story.append(Spacer(1, 10))

                story.append(Spacer(1, 6))

            doc.build(story)
            pdf_bytes = buffer.getvalue()

            with open(out_path, "wb") as f:
                f.write(pdf_bytes)

            return pdf_bytes, out_path

        except Exception as e:
            logger.warning(f"ReportLab PDF generation fallback triggered: {e}")
            # Fallback to PyMuPDF
            import pymupdf

            doc = pymupdf.open()
            page = doc.new_page()
            text_lines = [f"{title.upper()}", f"Author: {author}", "=" * 50, ""]
            for s in sections:
                if s.get("heading"):
                    text_lines.append(f"\n## {s['heading']}\n")
                if s.get("content"):
                    text_lines.append(s["content"])
                for b in s.get("bullets", []):
                    text_lines.append(f"- {b}")

            full_text = "\n".join(text_lines)
            page.insert_text((50, 50), full_text, fontsize=10)
            pdf_bytes = doc.tobytes()
            with open(out_path, "wb") as f:
                f.write(pdf_bytes)
            return pdf_bytes, out_path

    def generate_docx(
        self,
        title: str,
        sections: list[dict[str, Any]],
        filename: str | None = None,
        author: str = "C.O.P.P.E.R. AI",
        metadata: dict[str, Any] | None = None,
    ) -> tuple[bytes, Path]:
        """
        Generates a styled Microsoft Word (.docx) document using python-docx.
        """
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        fname = self._sanitize_filename(filename or title, "docx")
        out_path = self.output_dir / fname

        doc = docx.Document()

        # Set standard margins
        for sec in doc.sections:
            sec.top_margin = Inches(1.0)
            sec.bottom_margin = Inches(1.0)
            sec.left_margin = Inches(1.0)
            sec.right_margin = Inches(1.0)

        # Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42)  # #0f172a
        title_p.paragraph_format.space_after = Pt(4)

        # Subtitle / Author
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(f"Author: {author}  |  Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M UTC')}")
        sub_run.font.size = Pt(9.5)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 116, 139)
        sub_p.paragraph_format.space_after = Pt(18)

        # Sections
        for s in sections:
            heading = s.get("heading") or s.get("title")
            if heading:
                h_p = doc.add_heading(level=1)
                h_run = h_p.add_run(heading)
                h_run.font.size = Pt(14)
                h_run.font.bold = True
                h_run.font.color.rgb = RGBColor(200, 125, 85)  # Copper accent

            subheading = s.get("subheading")
            if subheading:
                sub_h = doc.add_heading(level=2)
                sub_h_run = sub_h.add_run(subheading)
                sub_h_run.font.size = Pt(11.5)
                sub_h_run.font.bold = True

            content = s.get("content") or s.get("text") or ""
            if content:
                for para in content.split("\n\n"):
                    if para.strip():
                        p = doc.add_paragraph()
                        p_run = p.add_run(para.strip())
                        p_run.font.size = Pt(10)
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.space_after = Pt(6)

            bullets = s.get("bullets") or s.get("items") or []
            if isinstance(bullets, list):
                for b in bullets:
                    b_p = doc.add_paragraph(style="List Bullet")
                    b_run = b_p.add_run(b)
                    b_run.font.size = Pt(10)

            # Table in docx
            table_data = s.get("table")
            if table_data and isinstance(table_data, dict):
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers or rows:
                    total_cols = len(headers) if headers else len(rows[0])
                    table = doc.add_table(rows=0, cols=total_cols)
                    table.style = "Table Grid"

                    if headers:
                        hdr_row = table.add_row().cells
                        for i, h in enumerate(headers):
                            hdr_row[i].text = str(h)
                            for paragraph in hdr_row[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(9.5)

                    for row in rows:
                        data_row = table.add_row().cells
                        for i, val in enumerate(row):
                            if i < len(data_row):
                                data_row[i].text = str(val)
                                for paragraph in data_row[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        with open(out_path, "wb") as f:
            f.write(docx_bytes)

        return docx_bytes, out_path

    def generate_markdown(
        self,
        title: str,
        sections: list[dict[str, Any]],
        filename: str | None = None,
        author: str = "C.O.P.P.E.R. AI",
        metadata: dict[str, Any] | None = None,
    ) -> tuple[str, Path]:
        """
        Generates a structured Markdown document with frontmatter and tables.
        """
        fname = self._sanitize_filename(filename or title, "md")
        out_path = self.output_dir / fname

        lines = [
            "---",
            f"title: \"{title}\"",
            f"author: \"{author}\"",
            f"date: \"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M UTC')}\"",
            f"generator: \"C.O.P.P.E.R. AI Document Engine\"",
            "---",
            "",
            f"# {title}",
            "",
            f"*Generated by {author} · {datetime.datetime.now().strftime('%B %d, %Y')}*",
            "",
            "---",
            "",
        ]

        for s in sections:
            heading = s.get("heading") or s.get("title")
            if heading:
                lines.append(f"## {heading}\n")

            subheading = s.get("subheading")
            if subheading:
                lines.append(f"### {subheading}\n")

            content = s.get("content") or s.get("text") or ""
            if content:
                lines.append(content.strip())
                lines.append("")

            bullets = s.get("bullets") or s.get("items") or []
            if isinstance(bullets, list) and bullets:
                for b in bullets:
                    lines.append(f"- {b}")
                lines.append("")

            table_data = s.get("table")
            if table_data and isinstance(table_data, dict):
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers:
                    lines.append("| " + " | ".join(str(h) for h in headers) + " |")
                    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                    for row in rows:
                        lines.append("| " + " | ".join(str(v) for v in row) + " |")
                    lines.append("")

        md_text = "\n".join(lines)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(md_text)

        return md_text, out_path

    def generate_html(
        self,
        title: str,
        sections: list[dict[str, Any]],
        filename: str | None = None,
        author: str = "C.O.P.P.E.R. AI",
        metadata: dict[str, Any] | None = None,
    ) -> tuple[str, Path]:
        """
        Generates a modern, responsive, standalone HTML document.
        """
        fname = self._sanitize_filename(filename or title, "html")
        out_path = self.output_dir / fname

        body_elements = []
        for s in sections:
            sec_html = ["<section class='doc-section'>"]
            if s.get("heading"):
                sec_html.append(f"<h2 class='section-title'>{s['heading']}</h2>")
            if s.get("subheading"):
                sec_html.append(f"<h3 class='section-subtitle'>{s['subheading']}</h3>")
            if s.get("content"):
                paras = [p.strip() for p in s["content"].split("\n\n") if p.strip()]
                for p in paras:
                    sec_html.append(f"<p>{p.replace('\n', '<br/>')}</p>")
            if s.get("bullets"):
                sec_html.append("<ul class='bullet-list'>")
                for b in s["bullets"]:
                    sec_html.append(f"<li>{b}</li>")
                sec_html.append("</ul>")

            table_data = s.get("table")
            if table_data and isinstance(table_data, dict):
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers or rows:
                    sec_html.append("<div class='table-container'><table>")
                    if headers:
                        sec_html.append("<thead><tr>" + "".join(f"<th>{h}</th>" for h in headers) + "</tr></thead>")
                    if rows:
                        sec_html.append("<tbody>")
                        for r in rows:
                            sec_html.append("<tr>" + "".join(f"<td>{cell}</td>" for cell in r) + "</tr>")
                        sec_html.append("</tbody>")
                    sec_html.append("</table></div>")

            sec_html.append("</section>")
            body_elements.append("\n".join(sec_html))

        html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>
    :root {{
      --bg: #0b0f17;
      --card-bg: #131b2e;
      --text-main: #f1f5f9;
      --text-muted: #94a3b8;
      --copper: #c87d55;
      --copper-light: #e09f7a;
      --border: #1e293b;
    }}
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
      background: var(--bg);
      color: var(--text-main);
      line-height: 1.6;
      padding: 40px 20px;
    }}
    .container {{
      max-width: 860px;
      margin: 0 auto;
      background: var(--card-bg);
      padding: 48px;
      border-radius: 16px;
      border: 1px solid var(--border);
      box-shadow: 0 20px 40px rgba(0,0,0,0.5);
    }}
    header {{
      border-bottom: 2px solid var(--copper);
      padding-bottom: 24px;
      margin-bottom: 32px;
    }}
    h1 {{
      font-size: 28px;
      font-weight: 800;
      color: #ffffff;
      margin-bottom: 8px;
      letter-spacing: -0.5px;
    }}
    .meta-bar {{
      font-size: 13px;
      color: var(--text-muted);
    }}
    .doc-section {{
      margin-bottom: 28px;
    }}
    .section-title {{
      font-size: 18px;
      font-weight: 700;
      color: var(--copper-light);
      margin-bottom: 12px;
      padding-bottom: 4px;
      border-bottom: 1px solid var(--border);
    }}
    .section-subtitle {{
      font-size: 14px;
      font-weight: 600;
      color: var(--text-main);
      margin-bottom: 8px;
    }}
    p {{
      margin-bottom: 14px;
      font-size: 15px;
      color: #cbd5e1;
    }}
    .bullet-list {{
      margin-left: 20px;
      margin-bottom: 16px;
    }}
    .bullet-list li {{
      margin-bottom: 6px;
      font-size: 14.5px;
      color: #cbd5e1;
    }}
    .table-container {{
      overflow-x: auto;
      margin: 16px 0;
      border-radius: 8px;
      border: 1px solid var(--border);
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 13.5px;
      text-align: left;
    }}
    th {{
      background: #1e293b;
      color: #ffffff;
      padding: 10px 14px;
      font-weight: 600;
      border-bottom: 1px solid var(--border);
    }}
    td {{
      padding: 10px 14px;
      border-bottom: 1px solid var(--border);
      color: #cbd5e1;
    }}
    tr:nth-child(even) td {{
      background: rgba(255,255,255,0.02);
    }}
    footer {{
      margin-top: 40px;
      padding-top: 16px;
      border-top: 1px solid var(--border);
      font-size: 12px;
      color: var(--text-muted);
      text-align: center;
    }}
  </style>
</head>
<body>
  <div class="container">
    <header>
      <h1>{title}</h1>
      <div class="meta-bar">Generated by <strong>{author}</strong> · {datetime.datetime.now().strftime('%B %d, %Y · %H:%M UTC')}</div>
    </header>
    <main>
      {"".join(body_elements)}
    </main>
    <footer>
      C.O.P.P.E.R. Autonomous Operating Environment · Document Services Engine
    </footer>
  </div>
</body>
</html>"""

        with open(out_path, "w", encoding="utf-8") as f:
            f.write(html_template)

        return html_template, out_path

    def generate_csv(
        self,
        headers: list[str],
        rows: list[list[Any]],
        filename: str | None = None,
    ) -> tuple[str, Path]:
        """
        Generates a clean CSV file.
        """
        fname = self._sanitize_filename(filename or "export", "csv")
        out_path = self.output_dir / fname

        output = io.StringIO()
        writer = csv.writer(output)
        if headers:
            writer.writerow(headers)
        for row in rows:
            writer.writerow(row)

        csv_text = output.getvalue()
        with open(out_path, "w", encoding="utf-8", newline="") as f:
            f.write(csv_text)

        return csv_text, out_path

    def generate_tsv(
        self,
        headers: list[str],
        rows: list[list[Any]],
        filename: str | None = None,
    ) -> tuple[str, Path]:
        """
        Generates a clean TSV file.
        """
        fname = self._sanitize_filename(filename or "export", "tsv")
        out_path = self.output_dir / fname

        output = io.StringIO()
        writer = csv.writer(output, delimiter="\t")
        if headers:
            writer.writerow(headers)
        for row in rows:
            writer.writerow(row)

        tsv_text = output.getvalue()
        with open(out_path, "w", encoding="utf-8", newline="") as f:
            f.write(tsv_text)

        return tsv_text, out_path

    def generate_json(
        self,
        data: Any,
        filename: str | None = None,
        indent: int = 2,
    ) -> tuple[str, Path]:
        """
        Generates a validated JSON document.
        """
        fname = self._sanitize_filename(filename or "data_export", "json")
        out_path = self.output_dir / fname

        json_text = json.dumps(data, indent=indent, ensure_ascii=False)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(json_text)

        return json_text, out_path

    def generate_text(
        self,
        title: str,
        sections: list[dict[str, Any]],
        filename: str | None = None,
        author: str = "C.O.P.P.E.R. AI",
    ) -> tuple[str, Path]:
        """
        Generates a clean formatted text document with ASCII headers.
        """
        fname = self._sanitize_filename(filename or title, "txt")
        out_path = self.output_dir / fname

        lines = [
            "=" * 70,
            f"  {title.upper()}",
            f"  Author: {author}  |  Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M UTC')}",
            "=" * 70,
            "",
        ]

        for s in sections:
            heading = s.get("heading") or s.get("title")
            if heading:
                lines.append(f"[{heading.upper()}]")
                lines.append("-" * len(heading))
            if s.get("content"):
                lines.append(s["content"].strip())
                lines.append("")
            for b in s.get("bullets", []):
                lines.append(f"  * {b}")
            if s.get("bullets"):
                lines.append("")

        txt = "\n".join(lines)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(txt)

        return txt, out_path

    async def create_document(
        self,
        format: str,
        title: str,
        sections: list[dict[str, Any]] | None = None,
        raw_content: str | None = None,
        headers: list[str] | None = None,
        rows: list[list[Any]] | None = None,
        data: Any = None,
        template_type: str = "general",
        filename: str | None = None,
        author: str = "C.O.P.P.E.R. AI",
        index_to_memory: bool = True,
    ) -> dict[str, Any]:
        """
        Universal high-level document generation entry point.
        Handles all document formats, files creation, metadata indexing, and returns standard response.
        """
        fmt = format.lower().strip().lstrip(".")
        if fmt not in ["pdf", "docx", "md", "html", "csv", "tsv", "json", "txt", "yaml"]:
            fmt = "pdf"

        # If raw content passed, convert to standard section format
        if not sections and raw_content:
            sections = [{"heading": "Overview", "content": raw_content}]
        elif not sections:
            sections = [{"heading": "Document Body", "content": f"Document: {title}"}]

        file_bytes: bytes = b""
        out_path: Path

        if fmt == "pdf":
            file_bytes, out_path = self.generate_pdf(title=title, sections=sections, filename=filename, author=author)
        elif fmt == "docx":
            file_bytes, out_path = self.generate_docx(title=title, sections=sections, filename=filename, author=author)
        elif fmt == "md":
            text_res, out_path = self.generate_markdown(title=title, sections=sections, filename=filename, author=author)
            file_bytes = text_res.encode("utf-8")
        elif fmt == "html":
            text_res, out_path = self.generate_html(title=title, sections=sections, filename=filename, author=author)
            file_bytes = text_res.encode("utf-8")
        elif fmt == "csv":
            text_res, out_path = self.generate_csv(headers=headers or ["Column 1", "Column 2"], rows=rows or [], filename=filename)
            file_bytes = text_res.encode("utf-8")
        elif fmt == "tsv":
            text_res, out_path = self.generate_tsv(headers=headers or ["Column 1", "Column 2"], rows=rows or [], filename=filename)
            file_bytes = text_res.encode("utf-8")
        elif fmt == "json":
            text_res, out_path = self.generate_json(data=data or {"title": title, "sections": sections}, filename=filename)
            file_bytes = text_res.encode("utf-8")
        else:
            text_res, out_path = self.generate_text(title=title, sections=sections, filename=filename, author=author)
            file_bytes = text_res.encode("utf-8")

        size_bytes = len(file_bytes)
        size_formatted = self.format_file_size(size_bytes)

        # Index to vector memory if requested
        indexed_chunks = 0
        if index_to_memory and raw_content:
            try:
                chunks = chunk_text(raw_content)
                for i, chunk in enumerate(chunks):
                    await memory_manager.save_document(
                        content=chunk,
                        source=out_path.name,
                        metadata={
                            "filename": out_path.name,
                            "extension": fmt,
                            "template": template_type,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                        },
                    )
                indexed_chunks = len(chunks)
            except Exception as e:
                logger.warning(f"Failed to vector-index generated doc '{out_path.name}': {e}")

        return {
            "status": "success",
            "filename": out_path.name,
            "filepath": str(out_path),
            "format": fmt,
            "title": title,
            "template_type": template_type,
            "size_bytes": size_bytes,
            "size_formatted": size_formatted,
            "indexed_chunks": indexed_chunks,
            "download_url": f"/documents/download/{out_path.name}",
            "created_at": datetime.datetime.now(datetime.timezone.utc).isoformat(),
        }

    def list_generated_documents(self) -> list[dict[str, Any]]:
        """
        Lists all generated documents in the output directory.
        """
        docs = []
        if not self.output_dir.exists():
            return []

        for f in sorted(self.output_dir.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
            if f.is_file():
                ext = f.suffix.lstrip(".").lower()
                stat = f.stat()
                docs.append(
                    {
                        "filename": f.name,
                        "filepath": str(f),
                        "extension": ext,
                        "category": self.SUPPORTED_EXTENSIONS.get(ext, f".{ext.upper()} File"),
                        "size_bytes": stat.st_size,
                        "size_formatted": self.format_file_size(stat.st_size),
                        "modified_at": datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "download_url": f"/documents/download/{f.name}",
                    }
                )
        return docs

    def get_document_file_path(self, filename: str) -> Path | None:
        """
        Safely resolves a file path within the documents directory preventing traversal attacks.
        """
        safe_name = Path(filename).name
        target = (self.output_dir / safe_name).resolve()
        if str(target).startswith(str(self.output_dir.resolve())) and target.is_file():
            return target
        return None

    # =========================================================================
    # DOCUMENT PARSING & EXTRACTION ENGINE (ORIGINAL METHODS)
    # =========================================================================

    async def parse_document(self, file_bytes: bytes, filename: str, index_to_memory: bool = True) -> dict[str, Any]:
        ext = Path(filename).suffix.lstrip(".").lower()
        size_bytes = len(file_bytes)
        size_formatted = self.format_file_size(size_bytes)
        file_category = self.SUPPORTED_EXTENSIONS.get(ext, f".{ext.upper()} File")

        pages: list[dict[str, Any]] = []
        full_text = ""
        structured_data: Any = None
        error_msg: str | None = None

        if ext == "pdf":
            pages, full_text, error_msg = self._parse_pdf(file_bytes)
        elif ext in ["csv", "tsv"]:
            delimiter = "," if ext == "csv" else "\t"
            pages, full_text, structured_data, error_msg = self._parse_delimited(file_bytes, delimiter)
        elif ext == "json":
            pages, full_text, structured_data, error_msg = self._parse_json(file_bytes)
        else:
            pages, full_text, error_msg = self._parse_text(file_bytes)

        lines = full_text.splitlines()
        words = full_text.split()
        word_count = len(words)
        char_count = len(full_text)
        estimated_tokens = max(1, int(char_count / 4))

        preview_lines = lines[:100]
        preview_text = "\n".join(preview_lines)
        if len(lines) > 100:
            preview_text += f"\n\n[... {len(lines) - 100} more lines in full document ({size_formatted}) ...]"

        indexed_chunks = 0
        if index_to_memory and full_text.strip():
            try:
                chunks = chunk_text(full_text)
                for i, chunk in enumerate(chunks):
                    await memory_manager.save_document(
                        content=chunk,
                        source=filename,
                        metadata={
                            "filename": filename,
                            "extension": ext,
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "word_count": word_count,
                        },
                    )
                indexed_chunks = len(chunks)
                logger.info(f"Indexed uploaded document '{filename}' ({indexed_chunks} chunks)")
            except Exception as e:
                logger.warning(f"Failed to vector-index document '{filename}': {e}")

        return {
            "filename": filename,
            "extension": ext,
            "category": file_category,
            "size_bytes": size_bytes,
            "size_formatted": size_formatted,
            "page_count": len(pages),
            "line_count": len(lines),
            "word_count": word_count,
            "char_count": char_count,
            "estimated_tokens": estimated_tokens,
            "indexed_chunks": indexed_chunks,
            "pages": pages,
            "full_text": full_text,
            "preview_text": preview_text,
            "structured_data": structured_data,
            "error": error_msg,
            "status": "success" if not error_msg else "partial",
        }

    def _parse_pdf(self, file_bytes: bytes) -> tuple[list[dict[str, Any]], str, str | None]:
        pages = []
        full_text_parts = []
        error = None

        try:
            import pypdf

            stream = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(stream)
            total_pages = len(reader.pages)

            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                clean_text = page_text.strip()
                page_words = len(clean_text.split())
                pages.append(
                    {
                        "page_number": idx + 1,
                        "text": clean_text,
                        "word_count": page_words,
                        "char_count": len(clean_text),
                    }
                )
                if clean_text:
                    full_text_parts.append(f"--- [Page {idx + 1} of {total_pages}] ---\n{clean_text}")

            full_text = "\n\n".join(full_text_parts)
            return pages, full_text, None
        except Exception as e:
            logger.warning(f"pypdf failed to parse PDF: {e}")
            error = str(e)

        try:
            import pymupdf

            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            total_pages = len(doc)
            pages = []
            full_text_parts = []

            for idx, page in enumerate(doc):
                text = page.get_text() or ""
                clean_text = text.strip()
                pages.append(
                    {
                        "page_number": idx + 1,
                        "text": clean_text,
                        "word_count": len(clean_text.split()),
                        "char_count": len(clean_text),
                    }
                )
                if clean_text:
                    full_text_parts.append(f"--- [Page {idx + 1} of {total_pages}] ---\n{clean_text}")

            full_text = "\n\n".join(full_text_parts)
            return pages, full_text, None
        except Exception as e2:
            logger.error(f"pymupdf also failed on PDF: {e2}")
            return [], "", f"Could not parse PDF: {error or e2}"

    def _parse_delimited(
        self, file_bytes: bytes, delimiter: str = ","
    ) -> tuple[list[dict[str, Any]], str, Any, str | None]:
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text), delimiter=delimiter)
            rows = list(reader)
            if not rows:
                return [], "", {"headers": [], "rows": [], "total_rows": 0}, None

            headers = rows[0]
            data_rows = rows[1:]
            preview_rows = data_rows[:50]

            structured = {
                "headers": headers,
                "preview_rows": preview_rows,
                "total_rows": len(data_rows),
                "column_count": len(headers),
            }

            pages = [
                {
                    "page_number": 1,
                    "text": text,
                    "word_count": len(text.split()),
                    "char_count": len(text),
                }
            ]
            return pages, text, structured, None
        except Exception as e:
            logger.error(f"CSV/TSV parse error: {e}")
            return [], "", None, str(e)

    def _parse_json(self, file_bytes: bytes) -> tuple[list[dict[str, Any]], str, Any, str | None]:
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            parsed_json = json.loads(text)
            formatted_text = json.dumps(parsed_json, indent=2)

            top_level_keys = list(parsed_json.keys()) if isinstance(parsed_json, dict) else []
            item_count = len(parsed_json) if isinstance(parsed_json, (dict, list)) else 1

            structured = {
                "is_array": isinstance(parsed_json, list),
                "is_object": isinstance(parsed_json, dict),
                "top_level_keys": top_level_keys[:30],
                "item_count": item_count,
            }

            pages = [
                {
                    "page_number": 1,
                    "text": formatted_text,
                    "word_count": len(formatted_text.split()),
                    "char_count": len(formatted_text),
                }
            ]
            return pages, formatted_text, structured, None
        except Exception as e:
            text = file_bytes.decode("utf-8", errors="replace")
            return (
                [{"page_number": 1, "text": text, "word_count": len(text.split()), "char_count": len(text)}],
                text,
                None,
                f"Invalid JSON: {e}",
            )

    def _parse_text(self, file_bytes: bytes) -> tuple[list[dict[str, Any]], str, str | None]:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode("latin-1")
            except Exception as e:
                return [], "", f"Could not decode text file: {e}"

        pages = [
            {
                "page_number": 1,
                "text": text,
                "word_count": len(text.split()),
                "char_count": len(text),
            }
        ]
        return pages, text, None


document_service = DocumentService()

